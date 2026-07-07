import os

def create_txt_resume(output_path: str):
    resume_text = """JANE DOE
Lead Backend Developer
Email: jane.doe@example.com | Phone: (123) 456-7890 | GitHub: github.com/janedoe

PROFESSIONAL SUMMARY
Highly experienced Backend Developer specializing in building scalable distributed systems, high-performance microservices, and robust data pipelines. Expert in Python, cloud technologies, and database optimization.

TECHNICAL SKILLS
* Languages: Python, Go, SQL, HTML/CSS, Bash
* Frameworks: FastAPI, Flask, gRPC, Django
* Databases & Caching: PostgreSQL, Redis, MongoDB, Elasticsearch
* Devops & Cloud: Docker, Kubernetes, AWS (EC2, S3, RDS), CI/CD (GitHub Actions)
* Data Engineering: Apache Spark, ETL Pipelines, Pandas

PROFESSIONAL EXPERIENCE
Lead Software Engineer | TechCorp (Jan 2023 - Present)
* Designed and implemented a high-throughput microservices architecture using Python, FastAPI, and gRPC, improving system latency by 35%.
* Built and optimized a distributed caching strategy using Redis, reducing core database read loads by 40%.
* Spearheaded migration of legacy monolithic apps to containerized environments using Docker and Kubernetes.
* Mentored 4 junior and mid-level engineers, establishing coding standards and CI/CD pipelines.

Software Engineer | DataInc (Jun 2020 - Jan 2023)
* Developed and maintained distributed ETL data pipelines using Apache Spark, processing over 10TB of data daily.
* Designed relational schemas and optimized queries in PostgreSQL, achieving a 20% database performance boost.
* Wrote clean, well-tested Python APIs (Flask) integrated with third-party payment gateways.

EDUCATION
BS in Computer Science | State University (2016 - 2020)
* GPA: 3.8/4.0
* Relevant coursework: Data Structures, Algorithms, Distributed Systems, Database Management Systems
"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(resume_text)
    print(f"TXT resume created at: {output_path}")

def create_docx_resume(output_path: str):
    try:
        import docx
    except ImportError:
        print("python-docx not installed. Skipping DOCX resume generation.")
        return

    doc = docx.Document()
    doc.add_heading('JANE DOE', 0)
    doc.add_paragraph('Lead Backend Developer\nEmail: jane.doe@example.com | Phone: (123) 456-7890')
    
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(
        "Highly experienced Backend Developer specializing in building scalable distributed systems, "
        "high-performance microservices, and robust data pipelines. Expert in Python, cloud technologies, "
        "and database optimization."
    )
    
    doc.add_heading('Technical Skills', level=1)
    doc.add_paragraph(
        "• Languages: Python, Go, SQL, Bash\n"
        "• Frameworks: FastAPI, Flask, gRPC\n"
        "• Databases & Caching: PostgreSQL, Redis, MongoDB\n"
        "• DevOps: Docker, Kubernetes, AWS, CI/CD"
    )
    
    doc.add_heading('Professional Experience', level=1)
    doc.add_heading('Lead Software Engineer | TechCorp (Jan 2023 - Present)', level=2)
    doc.add_paragraph(
        "• Designed and implemented a high-throughput microservices architecture using Python, FastAPI, and gRPC, improving system latency by 35%.\n"
        "• Built and optimized a distributed caching strategy using Redis, reducing core database read loads by 40%.\n"
        "• Spearheaded migration of legacy monolithic apps to containerized environments using Docker and Kubernetes."
    )
    
    doc.add_heading('Software Engineer | DataInc (Jun 2020 - Jan 2023)', level=2)
    doc.add_paragraph(
        "• Developed and maintained distributed ETL data pipelines using Apache Spark, processing over 10TB of data daily.\n"
        "• Designed relational schemas and optimized queries in PostgreSQL, achieving a 20% database performance boost."
    )
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph("BS in Computer Science | State University (2016 - 2020)")
    
    doc.save(output_path)
    print(f"DOCX resume created at: {output_path}")

def create_pdf_resume(output_path: str):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor
    except ImportError:
        print("reportlab not installed. Skipping PDF resume generation.")
        return

    doc = SimpleDocTemplate(output_path, pagesize=letter,
                            rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
    story = []
    styles = getSampleStyleSheet()
    
    # Custom colors
    primary_color = HexColor("#1E3A8A") # Navy
    text_color = HexColor("#374151") # Dark Grey
    
    # Title style
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        textColor=primary_color,
        spaceAfter=6
    )
    
    subtitle_style = ParagraphStyle(
        'DocSubTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=12,
        leading=14,
        textColor=text_color,
        spaceAfter=12
    )
    
    heading_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=18,
        textColor=primary_color,
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'BodyTextCustom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=text_color,
        spaceAfter=6
    )
    
    story.append(Paragraph("JANE DOE", title_style))
    story.append(Paragraph("Lead Backend Developer &nbsp;|&nbsp; jane.doe@example.com &nbsp;|&nbsp; (123) 456-7890", subtitle_style))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph("Professional Summary", heading_style))
    story.append(Paragraph(
        "Highly experienced Backend Developer specializing in building scalable distributed systems, "
        "high-performance microservices, and robust data pipelines. Expert in Python, cloud technologies, "
        "and database optimization.", body_style
    ))
    
    story.append(Paragraph("Technical Skills", heading_style))
    skills_text = (
        "<b>Languages:</b> Python, Go, SQL, Bash<br/>"
        "<b>Frameworks:</b> FastAPI, Flask, gRPC, Django<br/>"
        "<b>Databases & Caching:</b> PostgreSQL, Redis, MongoDB, Elasticsearch<br/>"
        "<b>Devops & Cloud:</b> Docker, Kubernetes, AWS, CI/CD"
    )
    story.append(Paragraph(skills_text, body_style))
    
    story.append(Paragraph("Professional Experience", heading_style))
    story.append(Paragraph("<b>Lead Software Engineer | TechCorp (Jan 2023 - Present)</b>", body_style))
    job1_text = (
        "• Designed and implemented a high-throughput microservices architecture using Python, FastAPI, and gRPC, improving system latency by 35%.<br/>"
        "• Built and optimized a distributed caching strategy using Redis, reducing database read loads by 40%.<br/>"
        "• Spearheaded migration of legacy monolithic apps to containerized environments using Docker and Kubernetes.<br/>"
        "• Mentored 4 junior engineers, establishing coding standards and CI/CD pipelines."
    )
    story.append(Paragraph(job1_text, body_style))
    
    story.append(Spacer(1, 5))
    story.append(Paragraph("<b>Software Engineer | DataInc (Jun 2020 - Jan 2023)</b>", body_style))
    job2_text = (
        "• Developed and maintained distributed ETL data pipelines using Apache Spark, processing over 10TB of data daily.<br/>"
        "• Designed relational schemas and optimized queries in PostgreSQL, achieving a 20% database performance boost."
    )
    story.append(Paragraph(job2_text, body_style))
    
    story.append(Paragraph("Education", heading_style))
    story.append(Paragraph("<b>BS in Computer Science</b> | State University (2016 - 2020)<br/>GPA: 3.8/4.0", body_style))
    
    doc.build(story)
    print(f"PDF resume created at: {output_path}")

if __name__ == "__main__":
    data_dir = os.path.dirname(os.path.abspath(__file__))
    
    create_txt_resume(os.path.join(data_dir, "sample_resume.txt"))
    create_docx_resume(os.path.join(data_dir, "sample_resume.docx"))
    create_pdf_resume(os.path.join(data_dir, "sample_resume.pdf"))
