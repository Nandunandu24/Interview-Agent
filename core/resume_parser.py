import os
import json
from .llm_client import ask_llm

def parse_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber is not installed. Please run pip install pdfplumber.")
        
    text_content = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
    return "\n".join(text_content)

def parse_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx is not installed. Please run pip install python-docx.")
        
    doc = docx.Document(file_path)
    text_content = []
    for para in doc.paragraphs:
        if para.text:
            text_content.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    text_content.append(cell.text)
    return "\n".join(text_content)

def parse_txt(file_path: str) -> str:
    """Extract text from a plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def parse_resume(file_path: str) -> str:
    """
    Parses a resume file based on its extension.
    Supports .pdf, .docx, and .txt.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Resume file not found at: {file_path}")
        
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    elif ext == ".txt":
        return parse_txt(file_path)
    else:
        # Fallback to text reading if possible
        try:
            return parse_txt(file_path)
        except Exception:
            raise ValueError(f"Unsupported resume file format: {ext}")

def infer_role_and_skills(resume_text: str) -> dict:
    """
    Uses the LLM client to analyze the candidate's resume text.
    Returns a dictionary with:
    {
      "inferred_role": "...",
      "skills": ["...", "..."],
      "summary": "..."
    }
    """
    prompt = f"""
You are an expert technical recruiter. Analyze the following candidate's resume text and extract the candidate's inferred target role, a list of their top 5-10 technical/professional skills, and a brief 2-sentence summary.

Return ONLY a valid JSON object matching the following schema:
{{
  "inferred_role": "String representing the best-guess target role based on experience and skills",
  "skills": ["Array", "of", "top", "skills", "identified"],
  "summary": "2-sentence professional summary of the candidate's profile"
}}

Do not include any markup, code blocks, or conversational text. Return only the JSON object.

Resume Text:
{resume_text}
"""
    response_text = ask_llm(prompt, expect_json=True)
    # Clean any markdown block formatting if LLM ignores instruction
    if "```json" in response_text:
        response_text = response_text.split("```json")[1].split("```")[0].strip()
    elif "```" in response_text:
        response_text = response_text.split("```")[1].split("```")[0].strip()
        
    data = json.loads(response_text.strip())
    return data
